#!/bin/python3

import os
import docx
from utils import (
    populate_callbacks,
    docx_tab_to_string_matrix,
    clean_match,
    docx_tab_to_string_matrix_for_api_def,
)
from docx.table import Table


def test_populate_callbacks():
    callbacks = {}
    input = [
        "The POST method creates a specific RNI event subscription.",
        "CellChangeSubscription",
        "RabEstSubscription",
        "RabModSubscription",
    ]
    multiple_remarks = True
    expected_output = {
        "{$request.body#/callbackUri}": {
            "post": {
                "summary": "Callback POST used to send a notification",
                "description": "Subscription notification",
                "operationId": "notificationPOST",
                "requestBody": {
                    "description": "Subscription notification",
                    "required": True,
                    "content": {
                        "application/json": {
                            "schema": {
                                "type": "object",
                                "properties": {
                                    "subscriptionNotification": {
                                        "oneOf": [
                                            {"$ref": "#/components/schemas/CellChangeSubscription"},
                                            {"$ref": "#/components/schemas/RabEstSubscription"},
                                            {"$ref": "#/components/schemas/RabModSubscription"},
                                        ]
                                    }
                                }
                            }
                        }
                    },
                },
                "responses": {
                    "204": {"description": "No content"},
                    "404": {"description": "Not found"},
                },
            }
        }
    }    
    actual_output = populate_callbacks(callbacks, multiple_remarks, input)
    sorted_actual_output = sort_dict_keys(actual_output)
    print(actual_output)
    assert actual_output == expected_output


def test_docx_tab_to_string_matrix():
    
    src_doc = "gs_mec028v020301p.docx"

    if os.getcwd()[-3:] == 'src':
        assert os.path.isfile("../mec_specs/" + src_doc)
        api_spec = docx.Document("../mec_specs/" + src_doc)
    else:
        assert os.path.isfile("mec_specs/" + src_doc)
        api_spec = docx.Document("mec_specs/" + src_doc)

    table = api_spec.tables[11]

    expected_output = [
        ["Attribute name", "Data type", "Cardinality", "Description"],
        [
            "notificationType",
            "String",
            "1",
            'Shall be set to "MeasurementReportNotification".',
        ],
    ]

    response = docx_tab_to_string_matrix(table)
    assert expected_output[:2] == response[:2]


def test_docx_tab_to_string_matrix_for_mec_010():
    
    src_doc = "gs_mec01002v030101p.docx"

    if os.getcwd()[-3:] == 'src':
        assert os.path.isfile("../mec_specs/" + src_doc)
        api_spec = docx.Document("../mec_specs/" + src_doc)
    else:
        assert os.path.isfile("mec_specs/" + src_doc)
        api_spec = docx.Document("mec_specs/" + src_doc)

    table = api_spec.tables[20]
    print(table)

    expected_output = [
        ["Attribute name", "Data type", "Cardinality", "Description"],
        [
            "serName",
            "String",
            "1",
            "The name of the service, for example, RNIS, LocationService, etc.",
        ],
    ]
    result = docx_tab_to_string_matrix(table, True)[:2]
    print(result)
    assert result == expected_output


def test_clean_match():
    """ """
    assert clean_match("A", "A")
    assert clean_match("6\tModel", "6 Model")
    assert not clean_match("6\tModel", "6.2 Information model")
    assert clean_match("6     Data model$", "6     Data model")
    assert clean_match("6     Data model", "6[ ]*Data model$")
    assert clean_match("6   Data model", "6[\s]*Data model$")


# ########## Test cases related to API Definitions implementation #################
# ################################################################################


def test_docx_tab_to_string_matrix_for_api_def():

    src_doc = "gs_mec028v020301p.docx"

    if os.getcwd()[-3:] == 'src':
        assert os.path.isfile("../mec_specs/" + src_doc)
        api_spec = docx.Document("../mec_specs/" + src_doc)
    else:
        assert os.path.isfile("mec_specs/" + src_doc)
        api_spec = docx.Document("mec_specs/" + src_doc)

    table = api_spec.tables[53]

    expected_output = [
        [
            "/measurements/\n{measurementConfigId}",
            "GET",
            "Retrieve information on an existing measurement configuration",
        ],
        [
            "/measurements/\n{measurementConfigId}",
            "PUT",
            "Modify an existing measurement configuration by sending a new data structure",
        ],
        [
            "/measurements/\n{measurementConfigId}",
            "DELETE",
            "Cancel an existing measurement configuration",
        ],
    ]
    assert docx_tab_to_string_matrix_for_api_def(table, False)[-3:] == expected_output


def sort_dict_keys(d: dict) -> dict:
    """
    Recursively sort dictionary keys alphabetically.

    Args:
        d (dict): A dictionary.

    Returns:
        dict: The dictionary with sorted keys.
    """
    sorted_dict = {}
    for k, v in sorted(d.items()):
        if isinstance(v, dict):
            sorted_dict[k] = sort_dict_keys(v)
        else:
            sorted_dict[k] = v
    return sorted_dict
